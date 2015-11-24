import os
from docx import Document
from docx.shared import Inches

def get(name):
	saveFile = os.getcwd()+"/"+name
	try:
		doc = Document()
		doc.add_heading("HEAD".decode('gbk'),0)
		doc.add_paragraph("HEAD CONTENT".decode('gbk'))
		doc.save(saveFile)
	except :
		print("error")

if __name__ == "__main__":
	get("test.docx")
	print("run")
